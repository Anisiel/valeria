import docx
import os
import re

# Percorso del file (controlla che sia nella cartella giusta)
word_file = 'testotesi.docx'
if not os.path.exists(word_file):
    word_file = 'documents/testotesi.docx'

if not os.path.exists(word_file):
    print("ERRORE: Non trovo il file testotesi.docx")
else:
    doc = docx.Document(word_file)
    content = "\n".join([p.text for p in doc.paragraphs])

    # Titolo che hai impostato tu (senza il 3)
    nuovo_titolo = "Analisi del tracciato della Via Valeria"
    
    # Cerchiamo il titolo ignorando maiuscole/minuscole e spazi extra
    pattern = re.escape(nuovo_titolo)
    match = re.search(pattern, content, re.IGNORECASE)

    if not match:
        # Piano B: cerchiamo almeno un punto di riferimento forte della sezione 3
        print("Titolo esatto non trovato, provo col Piano B (parole chiave)...")
        match = re.search(r'Da Tivoli a Vicovaro', content, re.IGNORECASE)

    if match:
        testo_integrale = content[match.start():]
        
        # Salvataggio
        output_dir = 'docs/frammenti_sparsi'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, '03_analisi_tracciato.md')

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# " + testo_integrale)
        
        print(f"SUCCESSO TOTALE! Sezione estratta in: {output_path}")
        print("-" * 30)
        print("ANTEPRIMA DEL TESTO ESTRATTO:")
        print(testo_integrale[:200] + "...")
    else:
        print("ERRORE CRITICO: Non trovo né il titolo né i capitoli interni.")
        print("Ecco i primi 10 titoli che Python vede nel tuo file:")
        for i, p in enumerate(doc.paragraphs[:100]):
            if len(p.text.strip()) > 5:
                print(f"- {p.text.strip()[:50]}")